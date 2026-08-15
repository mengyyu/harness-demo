"""Demo MCP Server — 提供演示工具"""

import io
from harness.mcp.manager import MCPTool


# ═══════════════════════════════════════════════════
# 真实文档解析
# ═══════════════════════════════════════════════════

def _extract_pdf(file_bytes: bytes) -> str:
    """从 PDF 提取文本"""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    """从 Word 文档提取文本"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_txt(file_bytes: bytes) -> str:
    """从纯文本提取"""
    for encoding in ("utf-8", "gbk", "gb18030", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("utf-8", errors="replace")


def parse_document_handler(file_path: str = "", file_name: str = "",
                           file_content: bytes = None, **kwargs) -> dict:
    """真实文档解析 — 支持 PDF / Word / 纯文本

    Args:
        file_path: 服务器本地文件路径（可选）
        file_name: 上传文件的原始文件名（用于判断格式）
        file_content: 上传文件的二进制内容（bytes）
    """
    # ── 优先解析上传的文件内容 ──
    if file_content:
        name = (file_name or file_path or "")
        ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""

        try:
            if ext == "pdf":
                raw_text = _extract_pdf(file_content)
                fmt = "pdf"
            elif ext in ("docx", "doc"):
                raw_text = _extract_docx(file_content)
                fmt = "docx"
            elif ext in ("txt", "md"):
                raw_text = _extract_txt(file_content)
                fmt = "txt"
            else:
                raw_text = _extract_txt(file_content)
                fmt = ext or "unknown"

            if raw_text.strip():
                return {
                    "success": True,
                    "file_name": file_name,
                    "format": fmt,
                    "raw_text": raw_text[:20000],
                    "source": "uploaded_file",
                }
        except Exception as e:
            return {
                "success": False,
                "file_name": file_name,
                "error": f"文档解析失败: {e}",
            }

    # ── 尝试读取服务器本地文件 ──
    if file_path and file_path != "demo_report.pdf":
        import os
        if os.path.exists(file_path):
            with open(file_path, "rb") as f:
                content = f.read()
            return parse_document_handler(
                file_name=os.path.basename(file_path), file_content=content,
            )

    # ── 兜底：无文件时返回演示文本 ──
    return {
        "success": True,
        "file_name": file_name or file_path or "demo_report.pdf",
        "format": "demo",
        "raw_text": """基金诊断报告
基金名称: 演示稳健增长混合型基金
基金代码: DEMO001
基金经理: 张三
诊断日期: 2024-06-30

业绩表现:
- 近1月收益率: 2.35%
- 近3月收益率: 5.82%
- 近1年收益率: 15.67%
- 年化波动率: 18.5%
- 夏普比率: 1.23

风险指标:
- 风险等级: R3
- 最大回撤: -15.8%

持仓分析:
- 行业分布: 消费28.5%, 科技22.3%, 金融18.7%
- 前三大重仓: 贵州茅台(8.5%), 宁德时代(6.2%), 招商银行(5.1%)
- 仓位集中度: 0.35

综合评分: 78分
""",
        "source": "demo",
    }





def get_account_analysis_handler(customer_id: str = "", **kwargs) -> dict:
    """模拟客户账户分析"""
    # Demo 数据（不同客户返回不同数据）
    customers = {
        "wang_001": {
            "customer_name": "王总",
            "customer_id": "wang_001",
            "analysis_date": "2024-06-30",
            "total_assets": 5000,
            "total_return": 12.5,
            "benchmark_return": 8.3,
            "risk_level": "R3",
            "risk_score": 62,
            "asset_allocation": {
                "股票型": 45,
                "混合型": 25,
                "债券型": 20,
                "货币型": 10,
            },
            "investment_horizon": "中长期",
            "risk_tolerance": "平衡型",
        },
        "li_001": {
            "customer_name": "李总",
            "customer_id": "li_001",
            "analysis_date": "2024-06-30",
            "total_assets": 8000,
            "total_return": 8.2,
            "benchmark_return": 8.3,
            "risk_level": "R2",
            "risk_score": 45,
            "asset_allocation": {
                "债券型": 50,
                "混合型": 30,
                "货币型": 15,
                "股票型": 5,
            },
            "investment_horizon": "短期",
            "risk_tolerance": "稳健型",
        },
        "zhang_001": {
            "customer_name": "张总",
            "customer_id": "zhang_001",
            "analysis_date": "2024-06-30",
            "total_assets": 12000,
            "total_return": 22.8,
            "benchmark_return": 8.3,
            "risk_level": "R4",
            "risk_score": 78,
            "asset_allocation": {
                "股票型": 65,
                "混合型": 20,
                "私募股权": 10,
                "债券型": 5,
            },
            "investment_horizon": "长期",
            "risk_tolerance": "进取型",
        },
    }

    return customers.get(customer_id, customers["wang_001"])


def get_holding_report_handler(customer_id: str = "", **kwargs) -> dict:
    """模拟私募持仓报告"""
    holdings = {
        "wang_001": {
            "customer_id": "wang_001",
            "report_date": "2024-06-30",
            "holdings": [
                {"name": "XX价值发现私募基金", "type": "股票多头", "nav": 1.85, "amount": 1500},
                {"name": "YY量化对冲1号", "type": "市场中性", "nav": 1.32, "amount": 800},
                {"name": "ZZ成长动力", "type": "股票多头", "nav": 2.15, "amount": 1200},
            ],
            "sector_allocation": {
                "消费": 28.5,
                "科技": 22.3,
                "金融": 18.7,
                "医药": 12.1,
                "新能源": 10.8,
                "其他": 7.6,
            },
            "concentration": {
                "top3_ratio": 0.42,
                "top5_ratio": 0.58,
                "herfindahl_index": 0.12,
            },
            "private_equity_products": [
                {
                    "name": "XX价值发现私募基金",
                    "manager": "李四",
                    "strategy": "价值投资",
                    "return_ytd": 15.3,
                    "return_1y": 22.1,
                    "max_drawdown": -12.5,
                    "sharpe": 1.45,
                    "aum": 50,
                },
                {
                    "name": "YY量化对冲1号",
                    "manager": "王五",
                    "strategy": "市场中性",
                    "return_ytd": 8.7,
                    "return_1y": 11.2,
                    "max_drawdown": -5.8,
                    "sharpe": 2.10,
                    "aum": 30,
                },
                {
                    "name": "ZZ成长动力",
                    "manager": "赵六",
                    "strategy": "成长股投资",
                    "return_ytd": 25.6,
                    "return_1y": 38.9,
                    "max_drawdown": -22.3,
                    "sharpe": 1.15,
                    "aum": 20,
                },
            ],
        },
        "li_001": {
            "customer_id": "li_001",
            "report_date": "2024-06-30",
            "holdings": [
                {"name": "AA稳健增利", "type": "债券型", "nav": 1.12, "amount": 3000},
                {"name": "BB固收+", "type": "混合型", "nav": 1.08, "amount": 2000},
            ],
            "sector_allocation": {"固定收益": 65, "货币市场": 20, "权益": 15},
            "concentration": {"top3_ratio": 0.85, "top5_ratio": 1.0, "herfindahl_index": 0.45},
            "private_equity_products": [
                {
                    "name": "AA稳健增利",
                    "manager": "钱七",
                    "strategy": "债券增强",
                    "return_ytd": 4.2,
                    "return_1y": 5.8,
                    "max_drawdown": -2.1,
                    "sharpe": 2.80,
                    "aum": 100,
                },
            ],
        },
        "zhang_001": {
            "customer_id": "zhang_001",
            "report_date": "2024-06-30",
            "holdings": [
                {"name": "CC科技先锋", "type": "股票多头", "nav": 3.25, "amount": 3500},
                {"name": "DD新兴产业", "type": "PE/VC", "nav": 1.95, "amount": 2500},
                {"name": "EE全球配置", "type": "QDII", "nav": 1.55, "amount": 2000},
            ],
            "sector_allocation": {
                "科技": 35.2,
                "新能源": 22.1,
                "消费": 15.8,
                "医药": 12.5,
                "金融科技": 8.2,
                "其他": 6.2,
            },
            "concentration": {
                "top3_ratio": 0.62,
                "top5_ratio": 0.78,
                "herfindahl_index": 0.18,
            },
            "private_equity_products": [
                {
                    "name": "CC科技先锋",
                    "manager": "孙八",
                    "strategy": "科技成长",
                    "return_ytd": 35.2,
                    "return_1y": 52.8,
                    "max_drawdown": -28.5,
                    "sharpe": 0.95,
                    "aum": 35,
                },
                {
                    "name": "DD新兴产业",
                    "manager": "周九",
                    "strategy": "PE/VC",
                    "return_ytd": 18.5,
                    "return_1y": 25.3,
                    "max_drawdown": -15.2,
                    "sharpe": 1.25,
                    "aum": 25,
                },
            ],
        },
    }

    return holdings.get(customer_id, holdings["wang_001"])


# ═══════════════════════════════════════════════════
# Tool 定义列表（供注册到 MCP Manager）
# ═══════════════════════════════════════════════════

DEMO_TOOLS = [
    MCPTool(
        name="parse_document",
        description="解析文档（PDF/Word/图片），提取文本和表格数据",
        parameters={
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "文件路径"},
            },
            "required": ["file_path"],
        },
        handler=parse_document_handler,
    ),
    MCPTool(
        name="get_account_analysis",
        description="获取客户账户智能分析结果",
        parameters={
            "type": "object",
            "properties": {
                "customer_id": {"type": "string", "description": "客户 ID"},
            },
            "required": ["customer_id"],
        },
        handler=get_account_analysis_handler,
    ),
    MCPTool(
        name="get_holding_report",
        description="获取客户私募持仓报告数据",
        parameters={
            "type": "object",
            "properties": {
                "customer_id": {"type": "string", "description": "客户 ID"},
            },
            "required": ["customer_id"],
        },
        handler=get_holding_report_handler,
    ),
]
